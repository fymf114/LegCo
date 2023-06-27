from bs4 import BeautifulSoup
from docx import Document
from io import BytesIO
from selenium import webdriver
from time import sleep
from tkinter import filedialog
from tkinter import Tk
from urllib3.exceptions import InsecureRequestWarning

import os
import pdfplumber
import re
import requests
import sys
import warnings


warnings.filterwarnings("ignore", category=InsecureRequestWarning)

start_phrase = input("Please input the Chinese full name of the desired LegCo Member.\n")
root = Tk()
root.withdraw()
output_dir = filedialog.askdirectory()
output_file_path = os.path.join(output_dir, 'speech.docx')

# For Chrome Driver
# chrome_driver_path = filedialog.askopenfilename(title="Select Chrome Driver", filetypes=(("All Files", "*.*"),))
# driver = webdriver.Chrome(chrome_driver_path)

driver = webdriver.Safari()

# Might have to replace with the URL when for future years
url = 'https://www.legco.gov.hk/tc/legco-business/council/council-meetings.html#schedule'
driver.get(url)
sleep(2)
links = []
soup = BeautifulSoup(driver.page_source, 'html.parser')
elements = soup.find_all('a')
for element in elements:
    href = element.get('href')
    if href and href.startswith('/tc/legco-business/council/hansard_rundown.html?f'):
        extracted_text = href.split('f')[1]
        # Find the pdf links of the official record of proceedings.
        # Would have to change the prefix for other years
        links.append('https://www.legco.gov.hk/yr2023/chinese/counmtg/floor/cm' + extracted_text + '-confirm-ec.pdf')
driver.quit()
# appending the link that is missing on the above LegCo page
links.append('https://www.legco.gov.hk/yr2023/chinese/counmtg/floor/cm20230524-confirm-ec.pdf')
start_phrase = start_phrase.strip() + '議員：'
end_phrases = ['議員：', '局長：', '主席：', '代理主席：']
document = Document()
current_dir = sys.path[0] if getattr(sys, 'frozen', False) else os.getcwd()
output_file_path = os.path.abspath(output_file_path)
is_speech_in_progress = False
speech_content = ""
pattern_to_be_removed = r'(立法會─\d+年\d+月\d+日)|(\d+LEGISLATIVECOUNCIL―\d+[A-Za-z]+\d+)|' \
                        r'(LEGISLATIVECOUNCIL―\d+[A-Za-z]+\d+)'

for link in links:
    try:
        response = requests.get(link, verify=False)
        response.raise_for_status()  # Raise an exception if the response is not successful

        pdf_bytes = BytesIO(response.content)
        with pdfplumber.open(pdf_bytes) as pdf:
            print("Processing: " + link)

            for page in pdf.pages:
                text = page.extract_text()
                text = re.sub(' ', '', text)  # Normalize spaces
                text = re.sub(pattern_to_be_removed, '', text)

                if start_phrase in text and not is_speech_in_progress:
                    text = text[text.index(start_phrase) + len(start_phrase):]
                    is_speech_in_progress = True

                if is_speech_in_progress:
                    for end_phrase in end_phrases:
                        if end_phrase in text:
                            text = text[:text.index(end_phrase)]
                            is_speech_in_progress = False
                            break

                    speech_content += text

                if not is_speech_in_progress and speech_content:
                    document.add_paragraph(link)
                    document.add_paragraph(speech_content.strip())
                    print(speech_content.strip())  # Print the desired paragraph
                    speech_content = ""

    except requests.exceptions.RequestException as e:
        print(f"Failed to download PDF from link: {link}")
        print(f"Error: {e}")
        continue

document.save(output_file_path)
